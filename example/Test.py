# -*- coding:utf-8 -*-

import os
import sys

from docx import Document
from docx.shared import Inches

sys.path.append(os.path.join(os.path.split(os.path.realpath(__file__))[0], ".."))

from DocProcessor import *

docENPath = "Test.en.docx"
docENModifyPath = "Test.en.md.docx"
docCNPath = "Test.cn.docx"
docCNModifyPath = "Test.cn.md.docx"

stringToFind = "China"
#stringToFind = "中国"
stringRegex = "China(.*?)d"
#stringRegex = "中(.*?)工"

docDebug = 1
parDebug = 1
runDebug = 1
apiDebug = 0

needSave = 0

isRegexTest = 0

docPath = docENPath
#docPath = docCNPath
docModifyPath = docENModifyPath
#docModifyPath = docCNModifyPath
def Test():
    # Open doc
    doc = Document(docPath)
    if docDebug:
        print("[[Path]] " + docPath)
        if apiDebug:
            print("[[Doc Objs]]")
            for obj in dir(doc):
                print("    " + obj)
            
    # Check each paragraph    
    lenParagraphs = len(doc.paragraphs)
    if parDebug:
        print("[[Paragraphs Len]] " + str(lenParagraphs))
    indexParagraph = 0
    for paragraph in doc.paragraphs:
        if parDebug:
            if indexParagraph == 0 and apiDebug:
                print("[[Paragraph Objs]]")
                for obj in dir(paragraph):
                    print("    " + obj)
            print("[[Paragraph " + str(indexParagraph) + "]] " + paragraph.text)

        # Check each run
        lenRuns = len(paragraph.runs)
        if runDebug:
            print("    [[Runs Len]] " + str(lenRuns))
        indexRun = 0
        for run in paragraph.runs:
            if runDebug:
                if indexParagraph == 0 and indexRun == 0 and apiDebug:
                    print("    [[Run Objs]]")
                    for obj in dir(run):
                        print("        " + obj)
                print("    [[Run " + str(indexRun) + "]] " + run.text)
            indexRun += 1

        indexParagraph += 1

    if needSave:
        doc.save(docModifyPath)

Test()

processor = Processor(docPath)
if processor.Open():
    print("Open")
    isFound = False
    if isRegexTest:
        isFound = processor.LocateRegexString(stringRegex)
    else:
        isFound = processor.LocateString(stringToFind)
    if isFound:
        print("Find")
        if processor.MarkString():
            print("Mark")
            if processor.Save():
                print("Saved")
            else:
                print("Could not save")
        else:
            print("Could not mark")
    else:
        print("Could not find string: " + stringToFind)
else:
    print("Could not to open: " + docPath)
