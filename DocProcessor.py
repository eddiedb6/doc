import os
import sys
import re

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.dml import MSO_COLOR_TYPE

sys.path.append(os.path.join(os.path.split(os.path.realpath(__file__))[0], "."))

from DocLocation import Location

class Processor:
    def __init__(self, path):
        self.__path = path
        self.__doc = None
        self.__locations = None
        self.__isRegex = False

    def Open(self):
        if not os.path.exists(self.__path):
            print("[[Doc]] Path does not exist: " + self.__path)
            return False
        self.__doc = Document(self.__path)
        return True

    def LocateRegexString(self, regex):
        if self.__doc is None:
            print("[[Doc]] Not ready for find yet")
            return False
        if regex is None or regex == "":
            print("[[Doc]] Invalid regex for find")
            return False
        self.__isRegex = True
        locateResult = self.__locateString(regex)
        self.__locations = locateResult[1]
        return locateResult[0]
    
    def LocateString(self, string):
        if self.__doc is None:
            print("[[Doc]] Not ready for locate yet")
            return False
        if string is None or string == "":
            print("[[Doc]] Invalid string for locate")
            return False
        self.__isRegex = False
        locateResult = self.__locateString(string)
        self.__locations = locateResult[1]
        return locateResult[0]

    def MarkString(self):
        if self.__doc is None:
            print("[[Doc]] Not ready for mark yet")
            return False
        if self.__locations is None:
            print("[[Doc]] No location for mark")
            return False
        locationsInParagraph = {}
        # Sort locations by paragraph and order by des
        for location in self.__locations:
            if location.ParagraphIndex in locationsInParagraph:
                locationsInParagraph[location.ParagraphIndex].insert(0, location)
            else:
                locationsInParagraph[location.ParagraphIndex] = [location]
        for paragraphIndex in locationsInParagraph:
            for location in locationsInParagraph[paragraphIndex]:
                if not self.__applyMark(location):
                    return False
        return True

    def Save(self):
        if self.__doc is None:
            print("[[Doc]] Not ready for save yet")
            return False
        self.__doc.save(self.__path)
        return True
        
    def __applyMark(self, location):
        paragraph = self.__doc.paragraphs[location.ParagraphIndex]
        runsCount = location.RunsCount

        print("[[Doc]] Apply Mark: " + str(location.ParagraphIndex))
        
        markRunIndex = -1
        preRunIndex = -1
        postRunOriginIndex = -1
        postRunMovedIndex = -1
        moveOffset = 0

        # Check head
        firstRun = location.GetRunIndex(0)
        firstStringRange = location.GetStringRange(firstRun)
        isFirstRunFromBeginning = firstStringRange[2]
        if isFirstRunFromBeginning:
            markRunIndex = firstRun
        else:
            markRunIndex = firstRun + 1
        preRunIndex = markRunIndex - 1
        postRunMovedIndex = markRunIndex + 1

        print("    markRunIndex:       " + str(markRunIndex))
        print("    preRunIndex:        " + str(preRunIndex))
        print("    postRunMovedIndex:  " + str(postRunMovedIndex))
                
        # Check end
        lastRun = location.GetRunIndex(runsCount -1)
        lastStringRange = location.GetStringRange(lastRun)
        isLastRunToEnd = lastStringRange[3]
        if isLastRunToEnd:
            postRunOriginIndex = lastRun + 1
        else:
            postRunOriginIndex = lastRun
        moveOffset = postRunMovedIndex - postRunOriginIndex

        print("    postRunOriginIndex: " + str(postRunOriginIndex))
        print("    moveOffset:         " + str(moveOffset))
        
        # Move
        if moveOffset == 1 or moveOffset == 2:
            # Add runs and move backward
            runAddCount = moveOffset
            while runAddCount > 0:
                paragraph.add_run()
                runAddCount -= 1
            index = len(paragraph.runs) - 1 - moveOffset
            while index >= postRunOriginIndex:
                self.__copyRun(paragraph.runs[index], paragraph.runs[index + moveOffset])
                index -= 1
        elif moveOffset > 2:
            print("[[Doc]] Should not move backward more than two")
            return False
        else:
            # Move forward
            offset = 0
            endIndex = len(paragraph.runs) - 1
            while postRunOriginIndex + offset <= endIndex:
                self.__copyRun(paragraph.runs[postRunOriginIndex + offset], paragraph.runs[postRunMovedIndex + offset])
                offset += 1
            # Reset useless run text in end
            while postRunMovedIndex + offset <= endIndex:
                paragraph.runs[postRunMovedIndex + offset].text = ""
                offset += 1

        # Handle head and end
        if not isFirstRunFromBeginning:
            paragraph.runs[preRunIndex].text = paragraph.runs[preRunIndex].text[:firstStringRange[0]]
        if not isLastRunToEnd:
            paragraph.runs[postRunMovedIndex].text = paragraph.runs[postRunMovedIndex].text[lastStringRange[1] + 1:]

        # Mark run
        markRun = paragraph.runs[markRunIndex]
        markRun.text = location.String
        markRun.bold = True
        markRun.italic = False
        markRun.underline = True
        markRunFont = markRun.font
        markRunFont.bold = True
        markRunFont.italic = False
        markRunFont.underline = True
        markRunFont.size = Pt(20)
        markRunFont.name = "Arial"
        markRunFontColor = markRunFont.color
        markRunFontColor.rgb = RGBColor(0xff, 0x00, 0x00)
        print(markRunFontColor.type)

        return True

    def __copyRun(self, runSrc, runDes):
        self.__copyFont(runSrc.font, runDes.font)
        runDes.bold =      runSrc.bold
        runDes.italic =    runSrc.italic
        runDes.style =     runSrc.style
        runDes.text =      runSrc.text
        runDes.underline = runSrc.underline

    def __copyFont(self, fontSrc, fontDes):
        fontDes.all_caps =        fontSrc.all_caps
        fontDes.bold =            fontSrc.bold
        fontDes.complex_script =  fontSrc.complex_script
        fontDes.cs_bold =         fontSrc.cs_bold
        fontDes.cs_italic =       fontSrc.cs_italic
        fontDes.double_strike =   fontSrc.double_strike
        fontDes.emboss =          fontSrc.emboss
        fontDes.hidden =          fontSrc.hidden
        fontDes.highlight_color = fontSrc.highlight_color
        fontDes.imprint =         fontSrc.imprint
        fontDes.italic =          fontSrc.italic
        fontDes.math =            fontSrc.math
        fontDes.name =            fontSrc.name
        fontDes.no_proof =        fontSrc.no_proof
        fontDes.outline =         fontSrc.outline
        fontDes.rtl =             fontSrc.rtl
        fontDes.shadow =          fontSrc.shadow
        fontDes.size =            fontSrc.size
        fontDes.small_caps =      fontSrc.small_caps
        fontDes.snap_to_grid =    fontSrc.snap_to_grid
        fontDes.spec_vanish =     fontSrc.spec_vanish
        fontDes.strike =          fontSrc.strike
        fontDes.subscript =       fontSrc.subscript
        fontDes.superscript =     fontSrc.superscript
        fontDes.underline =       fontSrc.underline
        fontDes.web_hidden =      fontSrc.web_hidden
        fontDes.color.theme_color = fontSrc.color.theme_color
        fontDes.color.rgb =       fontSrc.color.rgb

    def __locateStringInRun(self, paragraph, paragraphIndex, string, strBeginIndex):
        print("[[Doc]] String: " + string)
        strEndIndex = len(string) - 1 + strBeginIndex
        runIndex = 0
        runBeginIndex = 0
        runEndIndex = 0
        strPosInRun = {}
        def isStringNotMetYet():
            return True if runEndIndex < strBeginIndex else False
        def isStringMetAlready():
            return True if runBeginIndex > strEndIndex else False
        def isStringBeginHere():
            return True if runBeginIndex <= strBeginIndex and runEndIndex >= strBeginIndex else False
        def isStringEndHere():
            return True if runBeginIndex <= strEndIndex and runEndIndex >= strEndIndex else False
        def isRunInMiddleOfString():
            return True if runBeginIndex > strBeginIndex and runEndIndex < strEndIndex else False
        for run in paragraph.runs:
            runEndIndex = runBeginIndex + len(run.text) - 1
            if isStringNotMetYet():
                pass
            elif isStringMetAlready():
                print("[[Doc]] Warning: should not come to here")
                break
            elif isStringBeginHere():
                print("[[Doc]] Start in run: " + str(runIndex))
                if isStringEndHere():
                    print("[[Doc]] Also end in run: " + str(runIndex))
                    strPosInRun[runIndex] = [
                        strBeginIndex - runBeginIndex,
                        strEndIndex - runBeginIndex,
                        True if strBeginIndex == runBeginIndex else False,
                        True if strEndIndex == runEndIndex else False
                    ]
                    break
                else:
                    # Left part belongs to next run
                    strPosInRun[runIndex] = [
                        strBeginIndex - runBeginIndex,
                        runEndIndex - runBeginIndex,
                        True if strBeginIndex == runBeginIndex else False,
                        True
                    ]
            elif isRunInMiddleOfString():
                print("[[Doc]] This run is in middle of string: " + str(runIndex))
                strPosInRun[runIndex] = [
                    0,
                    runEndIndex - runBeginIndex,
                    True,
                    True
                ]
            elif isStringEndHere():
                print("[[Doc]] End in run: " + str(runIndex))
                strPosInRun[runIndex] = [
                    0,
                    strEndIndex - runBeginIndex,
                    True,
                    True if strEndIndex == runEndIndex else False
                ]
                break
            else:
                print("[[Doc]] Should not be here")
                return [False, None]
            runBeginIndex = runEndIndex + 1
            runIndex += 1
        if len(strPosInRun) <= 0:
            return [False, None]
        location = Location()
        location.String = string
        location.ParagraphIndex = paragraphIndex
        location.RunsCount = len(strPosInRun)
        count = 0
        print("[[Doc]] Location:")
        for runIndex in strPosInRun:
            location.SetRunIndex(count, runIndex)
            location.SetStringRange(runIndex, strPosInRun[runIndex][0], strPosInRun[runIndex][1], strPosInRun[runIndex][2], strPosInRun[runIndex][3])
            print("    run: " + str(runIndex))
            print("        start: " + str(strPosInRun[runIndex][0]))
            print("        end:   " + str(strPosInRun[runIndex][1]))
            print("        head:  " + str(strPosInRun[runIndex][2]))
            print("        tail:  " + str(strPosInRun[runIndex][3]))
            count += 1
        return [True, location]
        
    def __locateStringInParagraph(self, string, paragraph, paragraphIndex):
        locations = []
        beginIndex = 0
        strToLocate = None
        
        while True:
            resultStartIndex = -1

            if self.__isRegex:
                prog = re.compile(string)
                result = prog.search(paragraph.text, beginIndex)
                if not result:
                    break
                resultStartIndex = result.start(1)
                strToLocate = result.group(1)
            else:
                result = paragraph.text.find(string, beginIndex)
                if result < 0:
                    break
                resultStartIndex = result
                strToLocate = string
                
            # One found in paragraph
            print("[[Doc]] Find string in paragraph: " + str(paragraphIndex))
            locatedResult = self.__locateStringInRun(paragraph, paragraphIndex, strToLocate, resultStartIndex)
            if locatedResult[0]:
                locations.append(locatedResult[1])
            beginIndex = resultStartIndex + len(strToLocate)
            
        if len(locations) > 0:
            return [True, locations]
        return [False, None]

    def __locateString(self, string):
        locations = []
        paragraphIndex = 0
        for paragraph in self.__doc.paragraphs:
            locateResult = self.__locateStringInParagraph(string, paragraph, paragraphIndex)
            if locateResult[0]:
                for location in locateResult[1]:
                    locations.append(location)
            paragraphIndex += 1
        if len(locations) > 0:
            return [True, locations]
        return [False, None]
