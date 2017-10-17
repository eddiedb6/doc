from docx import Document
from docx.shared import Inches

docPath = "test.docx"
docModifyPath = "test.modify.docx"
debug = 1

# Open doc
doc = Document(docPath)
if debug:
    print("[[Path]] " + docPath)
    print("[[Doc Objs]]")
    for obj in dir(doc):
        print("    " + obj)
    
lenParagraphs = len(doc.paragraphs)
if debug:
    print("[[Paragraphs Len]] " + str(lenParagraphs))

# Check each paragraph
indexParagraph = 0
for paragraph in doc.paragraphs:
    if debug:
        if indexParagraph == 0:
            print("[[Paragraph Objs]]")
            for obj in dir(paragraph):
                print("    " + obj)
        print("[[Paragraph " + str(indexParagraph) + "]] " + paragraph.text)

    # Find out match words in paragraph
    
        
    lenRuns = len(paragraph.runs)
    if debug:
        print("    [[Runs Len]] " + str(lenRuns))

    # Check each run
    indexRun = 0
    for run in paragraph.runs:
        if debug:
            if indexParagraph == 0 and indexRun == 0:
                print("    [[Run Objs]]")
                for obj in dir(run):
                    print("        " + obj)
            print("    [[Run " + str(indexRun) + "]] " + run.text)
            
        indexRun += 1

    indexParagraph += 1

doc.save(docModifyPath)
